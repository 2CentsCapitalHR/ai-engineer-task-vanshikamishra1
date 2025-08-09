
from dataclasses import dataclass, field
from typing import Optional, Dict, Any, List
from pathlib import Path
import logging
import tempfile
import re
import io

# Third-party libs (ensure in requirements.txt)
import fitz  # PyMuPDF
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)
logger.setLevel(logging.INFO)


# -------------------------
# Data models
# -------------------------
@dataclass
class DocumentMetadata:
    filename: str
    document_type: str = "unknown"
    word_count: int = 0
    confidence_score: float = 1.0
    extra: Dict[str, Any] = field(default_factory=dict)


@dataclass
class DocumentContent:
    metadata: DocumentMetadata
    raw_text: str


# -------------------------
# Parser implementation
# -------------------------
class ADGMDocumentParser:
    """
    High-level parser that accepts uploaded files (file-like or path) and returns DocumentContent.
    Use parse_uploaded_file(uploaded_file) from Streamlit or other upload handlers.
    """

    # Heuristic keywords per ADGM document types (expandable)
    TYPE_KEYWORDS: Dict[str, List[str]] = {
        "articles_of_association": [
            "articles of association", "aoa", "articles", "constitution of the company"
        ],
        "memorandum_of_association": [
            "memorandum of association", "moa", "memorandum", "object clause"
        ],
        "board_resolution": [
            "board resolution", "resolved that", "meeting of the board", "quorum"
        ],
        "incorporation_application": [
            "incorporation application", "application for incorporation", "company incorporation"
        ],
        "ubo_declaration": [
            "ubo", "ultimate beneficial owner", "beneficial owner declaration"
        ],
        "register_of_members": [
            "register of members", "register of directors", "register of members and directors"
        ],
        "employment_contract": [
            "employment contract", "terms of employment", "employee", "employer"
        ],
        "license_application": [
            "license application", "licence application", "licensing", "fsra"
        ]
    }

    SUPPORTED_EXTS = {".docx", ".pdf", ".txt"}

    def __init__(self, min_word_confidence: int = 30):
        """
        :param min_word_confidence: heuristic minimum words to consider extraction confident
        """
        self.min_word_confidence = min_word_confidence

    # -------------------------
    # Public API
    # -------------------------
    def parse_uploaded_file(self, uploaded_file) -> DocumentContent:
        """
        Main entry point for uploaded files (Streamlit's UploadedFile or any file-like with .read() and .name).
        Returns DocumentContent dataclass.
        """
        try:
            filename = getattr(uploaded_file, "name", None) or getattr(uploaded_file, "filename", None)
            if not filename:
                raise ValueError("Uploaded file must have a filename attribute")

            name = Path(filename).name
            ext = Path(name).suffix.lower()

            if ext not in self.SUPPORTED_EXTS:
                raise ValueError(f"Unsupported file extension: {ext}")

            # Read bytes (works for streamlit UploadedFile and file-like objects)
            if hasattr(uploaded_file, "getbuffer"):
                data = uploaded_file.getbuffer().tobytes()
            else:
                # fallback: read()
                uploaded_file.seek(0)
                data = uploaded_file.read()

            # Save always in a known folder (avoids OneDrive/temp permission issues)
            # Save always in a known folder (avoids OneDrive/temp permission issues)
            tmp_dir = Path("tmp_uploads")
            tmp_dir.mkdir(exist_ok=True)
            tmp_path = tmp_dir / name
            with open(tmp_path, "wb") as wf:
                wf.write(data)



            # Dispatch by extension
            if ext == ".pdf":
                text = self._extract_pdf_text(str(tmp_path))
            elif ext == ".docx":
                text = self._extract_docx_text(str(tmp_path))
            elif ext == ".txt":
                text = self._extract_txt_text(str(tmp_path))
            else:
                text = ""

            # Clean up temp file
            try:
                tmp_path.unlink(missing_ok=True)
            except Exception:
                logger.debug("Temporary file cleanup failed: %s", tmp_path)

            # Normalize extracted text
            clean_text = self._clean_text(text)

            # Derive metadata
            word_count = self._word_count(clean_text)
            doc_type = self._detect_document_type(name, clean_text)
            confidence = self._estimate_confidence(word_count)

            metadata = DocumentMetadata(
                filename=name,
                document_type=doc_type,
                word_count=word_count,
                confidence_score=confidence,
                extra={}
            )

            logger.info("Parsed file: %s | type=%s | words=%d | conf=%.2f", name, doc_type, word_count, confidence)
            return DocumentContent(metadata=metadata, raw_text=clean_text)

        except Exception as e:
            logger.exception("Failed to parse uploaded file: %s", e)
            # Return an 'empty' DocumentContent with error metadata so pipeline can handle it gracefully
            metadata = DocumentMetadata(
                filename=getattr(uploaded_file, "name", "unknown"),
                document_type="unknown",
                word_count=0,
                confidence_score=0.0,
                extra={"error": str(e)}
            )
            return DocumentContent(metadata=metadata, raw_text="")

    # -------------------------
    # Extraction primitives
    # -------------------------
    def _extract_pdf_text(self, file_path: str) -> str:
        """
        Extracts text from PDF using PyMuPDF (fitz). Tries to be robust to pages with images,
        but does not perform OCR here.
        """
        try:
            doc = fitz.open(file_path)
            parts = []
            for page in doc:
                try:
                    txt = page.get_text("text")
                    if txt:
                        parts.append(txt)
                except Exception as e:
                    logger.debug("Page text extraction failed: %s", e)
            doc.close()
            return "\n".join(parts)
        except Exception as e:
            logger.exception("PDF extraction failed for %s: %s", file_path, e)
            return ""

    def _extract_docx_text(self, file_path: str) -> str:
        """
        Extracts text from a DOCX using python-docx.
        """
        try:
            doc = DocxDocument(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
            # Also try tables (basic handling) for textual cells
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text and cell.text.strip():
                            paragraphs.append(cell.text)
            return "\n".join(paragraphs)
        except Exception as e:
            logger.exception("DOCX extraction failed for %s: %s", file_path, e)
            return ""

    def _extract_txt_text(self, file_path: str) -> str:
        try:
            return Path(file_path).read_text(encoding="utf-8", errors="ignore")
        except Exception:
            return ""

    # -------------------------
    # Heuristics / Utilities
    # -------------------------
    def _clean_text(self, text: Optional[str]) -> str:
        if not text:
            return ""
        # Normalize whitespace and remove repeated header/footer lines that are likely artifacts
        # Collapse multiple newlines
        t = re.sub(r"\r\n?", "\n", text)
        t = re.sub(r"[ \t]+", " ", t)
        t = re.sub(r"\n{2,}", "\n\n", t)
        return t.strip()

    def _word_count(self, text: str) -> int:
        if not text:
            return 0
        return len(re.findall(r"\w+", text))

    def _estimate_confidence(self, word_count: int) -> float:
        """
        Very simple heuristic:
        - > min_word_confidence words => confidence 0.85 - 0.99 (scaled)
        - otherwise low confidence
        """
        if word_count == 0:
            return 0.0
        if word_count >= self.min_word_confidence:
            # scale between 0.85 and 0.99
            # cap at 2000 words for scaling stability
            wc = min(word_count, 2000)
            conf = 0.85 + 0.14 * (wc - self.min_word_confidence) / (2000 - self.min_word_confidence)
            return round(min(conf, 0.99), 2)
        # small docs may still be valid but set modest confidence
        return round(max(0.35, (word_count / self.min_word_confidence) * 0.6), 2)

    def _detect_document_type(self, filename: str, text: str) -> str:
        """
        Heuristic to guess the document type using filename + content.
        Returns one of the TYPE_KEYWORDS keys or 'unknown'.
        """
        name = filename.lower()
        lower_text = (text or "").lower()

        # check filename first
        for dtype, keys in self.TYPE_KEYWORDS.items():
            for k in keys:
                if k in name:
                    return dtype

        # fallback to scanning document text (count matches)
        scores = {k: 0 for k in self.TYPE_KEYWORDS}
        for dtype, keys in self.TYPE_KEYWORDS.items():
            for k in keys:
                if k in lower_text:
                    scores[dtype] += 1

        # choose best match if it has at least one hit
        best = max(scores.items(), key=lambda x: x[1])
        if best[1] > 0:
            return best[0]

        # final fallback - short heuristics based on distinct tokens
        if "agreement" in lower_text or "party" in lower_text:
            return "commercial_agreement"
        if "employee" in lower_text and "salary" in lower_text:
            return "employment_contract"

        return "unknown"
