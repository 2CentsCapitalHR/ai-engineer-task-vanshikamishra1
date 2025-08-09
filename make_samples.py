

from docx import Document
from pathlib import Path

SAMPLES_DIR = Path("samples")
SAMPLES_DIR.mkdir(exist_ok=True)

def make_aoa():
    doc = Document()
    doc.add_heading("Articles of Association", level=1)
    doc.add_paragraph("Company Name: Example ADGM Ltd")
    doc.add_paragraph("Registered Office: 123 Business Bay")
    # purposely omit "ADGM Courts" jurisdiction to trigger a flag
    doc.add_paragraph("Objects of the Company: To provide tech services.")
    doc.add_paragraph("Share Capital: 1,000,000 AED")
    doc.add_paragraph("Directors Powers: Board can appoint CEO.")
    doc.add_paragraph("\n") 
    doc.add_paragraph("Signature: ____________________")
    p = SAMPLES_DIR / "Articles_of_Association_sample.docx"
    doc.save(p)
    return p

def make_moa():
    doc = Document()
    doc.add_heading("Memorandum of Association", level=1)
    doc.add_paragraph("Company Name: Example ADGM Ltd")
    doc.add_paragraph("Registered Office Address: 123 Business Bay")
    doc.add_paragraph("Objects Clause: Provide consulting services.")
    # Missing subscriber details intentionally to create missing clause flag
    doc.add_paragraph("Authorized Share Capital: 1,000,000 AED")
    doc.add_paragraph("\n")
    p = SAMPLES_DIR / "Memorandum_of_Association_sample.docx"
    doc.save(p)
    return p

if __name__ == "__main__":
    print("Creating sample documents...")
    aoa = make_aoa()
    moa = make_moa()
    print("Created:", aoa, moa)
