
import json
from pathlib import Path
from datetime import datetime
from typing import Dict, Any, Optional
import logging

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

logger = logging.getLogger(__name__)
logger.setLevel(logging.INFO)


class ReportGenerator:
    def __init__(self, output_dir: str = "reports"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def generate_reports(self, aggregate_report: Dict[str, Any], prefix: Optional[str] = None) -> Dict[str, str]:
        """
        Generate report files for the provided aggregate_report dictionary.

        Returns dict with paths:
            {
                "json": "<path>",
                "docx": "<path>",
                "html": "<path or ''>",
                "pdf": "<path or ''>"
            }
        """
        timestamp = datetime.utcnow().strftime("%Y%m%dT%H%M%SZ")
        name_prefix = prefix or f"adgm_report_{timestamp}"
        json_path = self.output_dir / f"{name_prefix}.json"
        docx_path = self.output_dir / f"{name_prefix}.docx"
        html_path = self.output_dir / f"{name_prefix}.html"
        pdf_path = self.output_dir / f"{name_prefix}.pdf"

        # 1) JSON (exact required output)
        try:
            with open(json_path, "w", encoding="utf-8") as jf:
                json.dump(aggregate_report, jf, ensure_ascii=False, indent=2)
            logger.info("JSON report written: %s", json_path)
        except Exception as e:
            logger.exception("Failed to write JSON report: %s", e)
            raise

        # 2) DOCX human-readable report
        try:
            self._write_docx_report(aggregate_report, docx_path)
            logger.info("DOCX report written: %s", docx_path)
        except Exception as e:
            logger.exception("Failed to write DOCX report: %s", e)
            # continue; docx optional for JSON requirement

        # 3) Optional HTML (simple) and PDF via pdfkit if available
        try:
            html_content = self._render_html(aggregate_report)
            html_path.write_text(html_content, encoding="utf-8")
            logger.info("HTML report written: %s", html_path)
            # Try PDF via pdfkit (best-effort)
            try:
                import pdfkit
                pdfkit.from_file(str(html_path), str(pdf_path))
                logger.info("PDF report written: %s", pdf_path)
            except Exception:
                logger.info("pdfkit not available or conversion failed; skip PDF export.")
                pdf_path = Path("")
        except Exception:
            html_path = Path("")
            pdf_path = Path("")

        return {
            "json": str(json_path),
            "docx": str(docx_path) if docx_path.exists() else "",
            "html": str(html_path) if html_path and html_path.exists() else "",
            "pdf": str(pdf_path) if pdf_path and pdf_path.exists() else ""
        }

    def _write_docx_report(self, agg: Dict[str, Any], out_path: Path) -> None:
        doc = Document()
        # Title
        title = doc.add_heading("ADGM Corporate Agent — Compliance Report", level=1)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        meta = doc.add_paragraph()
        meta.add_run(f"Generated at: {datetime.utcnow().isoformat()}").italic = True
        meta.add_run("\n").bold = False
        meta.add_run(f"Process detected: {agg.get('process', 'Unknown')}\n")
        meta.add_run(f"Documents uploaded: {agg.get('documents_uploaded', 0)}\n")
        meta.add_run(f"Required documents (count): {agg.get('required_documents', 0)}\n")
        missing = agg.get("missing_documents") or []
        if missing:
            p = doc.add_paragraph()
            p.add_run("Missing documents:").bold = True
            for m in missing:
                doc.add_paragraph(f" - {m}", style="List Bullet")
        else:
            doc.add_paragraph("Missing documents: None")

        doc.add_paragraph("\nOverall issues summary:", style="Intense Quote")
        issues = agg.get("issues_found", [])
        severity_counts = {"critical": 0, "high": 0, "medium": 0, "low": 0}
        for i in issues:
            s = (i.get("severity") or "low").lower()
            if s in severity_counts:
                severity_counts[s] += 1

        for k, v in severity_counts.items():
            doc.add_paragraph(f"{k.title()}: {v}")

        doc.add_paragraph("\nPer-document summaries:\n")
        per_docs = agg.get("per_document_reports", [])
        for pr in per_docs:
            doc.add_heading(pr.get("document_name", "Unknown"), level=3)
            doc.add_paragraph(f"Type: {pr.get('document_type', 'unknown')}")
            doc.add_paragraph(f"Overall score: {pr.get('overall_score', 'N/A')}")
            doc.add_paragraph("Recommendations:")
            for r in pr.get("recommendations", []):
                doc.add_paragraph(f"- {r}", style="List Bullet")
            doc.add_paragraph("Missing clauses:")
            for mc in pr.get("missing_clauses", []):
                doc.add_paragraph(f"- {mc}", style="List Bullet")
            doc.add_paragraph("Validation summary:")
            vs = pr.get("validation_summary", {})
            for k, v in vs.items():
                doc.add_paragraph(f"  {k}: {v}")

        # Add top 20 issues as appendix
        doc.add_page_break()
        doc.add_heading("Issues (top 50)", level=2)
        count = 0
        for issue in issues:
            if count >= 50:
                break
            count += 1
            doc.add_paragraph(f"[{issue.get('severity', '').upper()}] {issue.get('document')} — {issue.get('issue')}")
            if issue.get("suggestion"):
                doc.add_paragraph(f"  Suggestion: {issue.get('suggestion')}")
            if issue.get("adgm_reference"):
                doc.add_paragraph(f"  ADGM reference: {issue.get('adgm_reference')}")

        # Styling tweaks (font sizes)
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)

        doc.save(out_path)

    def _render_html(self, agg: Dict[str, Any]) -> str:
        # Simple, clean HTML representation
        html = [
            "<!doctype html><html><head><meta charset='utf-8'/>",
            "<title>ADGM Compliance Report</title>",
            "<style>",
            "body{font-family:Arial,Helvetica,sans-serif;margin:24px;color:#111}",
            "h1{color:#0b3d91}",
            ".meta{color:#444;margin-bottom:12px}",
            ".section{margin-top:18px}",
            ".issue{margin-bottom:8px;padding:8px;border-left:4px solid #ddd}",
            ".critical{border-color:#c00}",
            ".high{border-color:#f60}",
            ".medium{border-color:#f90}",
            ".low{border-color:#0a0}"
            "</style></head><body>"
        ]
        html.append("<h1>ADGM Corporate Agent — Compliance Report</h1>")
        html.append(f"<div class='meta'>Generated at: {datetime.utcnow().isoformat()}</div>")
        html.append(f"<div class='section'><strong>Process:</strong> {agg.get('process')}</div>")
        html.append(f"<div class='section'><strong>Documents uploaded:</strong> {agg.get('documents_uploaded')} (required: {agg.get('required_documents')})</div>")

        missing = agg.get("missing_documents", [])
        if missing:
            html.append("<div class='section'><strong>Missing documents:</strong><ul>")
            for m in missing:
                html.append(f"<li>{m}</li>")
            html.append("</ul></div>")
        else:
            html.append("<div class='section'><strong>Missing documents:</strong> None</div>")

        html.append("<div class='section'><h2>Top Issues</h2>")
        for issue in agg.get("issues_found", [])[:200]:
            sev = (issue.get("severity") or "low").lower()
            html.append(f"<div class='issue {sev}'><strong>[{sev.upper()}]</strong> {issue.get('document')} — {issue.get('issue')}<br>")
            if issue.get("suggestion"):
                html.append(f"<em>Suggestion:</em> {issue.get('suggestion')}<br>")
            if issue.get("adgm_reference"):
                html.append(f"<em>Reference:</em> {issue.get('adgm_reference')}")
            html.append("</div>")
        html.append("</div></body></html>")
        return "".join(html)
 
