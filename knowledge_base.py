from __future__ import annotations
import json
import logging
from pathlib import Path
from typing import List, Dict, Any, Optional
import os

import numpy as np
from sentence_transformers import SentenceTransformer
import faiss
import fitz  # PyMuPDF
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)
logger.setLevel(logging.INFO)


# ---------------------------
# Helpers
# ---------------------------
def _chunk_text(text: str, chunk_size: int = 300, overlap: int = 50) -> List[str]:
    if not text:
        return []
    words = text.split()
    if len(words) <= chunk_size:
        return [" ".join(words)]
    chunks = []
    i = 0
    while i < len(words):
        chunk = words[i:i + chunk_size]
        chunks.append(" ".join(chunk))
        i += chunk_size - overlap
    return chunks


def _extract_text_from_pdf(path: Path) -> str:
    try:
        doc = fitz.open(str(path))
        pages = [p.get_text("text") for p in doc]
        doc.close()
        return "\n".join(pages)
    except Exception as e:
        logger.exception("Failed to extract PDF text from %s: %s", path, e)
        return ""


def _extract_text_from_docx(path: Path) -> str:
    try:
        doc = DocxDocument(str(path))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)
        return "\n".join(parts)
    except Exception as e:
        logger.exception("Failed to extract DOCX text from %s: %s", path, e)
        return ""


def _read_txt(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8", errors="ignore")
    except Exception as e:
        logger.exception("Failed to read TXT %s: %s", path, e)
        return ""


# ---------------------------
# ADGMKnowledgeBase
# ---------------------------
class ADGMKnowledgeBase:
    def __init__(
        self,
        db_dir: str = "kb_store",
        embedding_model_name: str = "all-MiniLM-L6-v2",
        chunk_size: int = 300,
        chunk_overlap: int = 50,
        db_path: Optional[str] = None  # backward compatibility
    ):
        if db_path:  # backward compatibility
            db_dir = Path(db_path).parent

        self.db_dir = Path(db_dir)
        self.db_dir.mkdir(parents=True, exist_ok=True)
        self.index_path = self.db_dir / "faiss_index.ivf"
        self.passages_path = self.db_dir / "passages.json"
        self.meta_path = self.db_dir / "meta.json"

        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

        self.model = SentenceTransformer(embedding_model_name)
        self.dim = self.model.get_sentence_embedding_dimension()

        self.index: Optional[faiss.Index] = None
        self.passages: List[Dict[str, Any]] = []

        if self.index_path.exists() and self.passages_path.exists():
            try:
                self._load()
                logger.info("Loaded existing KB with %d passages.", len(self.passages))
            except Exception as e:
                logger.exception("Failed to load existing KB: %s", e)

    # -------------------------
    # Persistence
    # -------------------------
    def _save(self) -> None:
        if self.index is None:
            return
        faiss.write_index(self.index, str(self.index_path))
        with open(self.passages_path, "w", encoding="utf-8") as f:
            json.dump(self.passages, f, ensure_ascii=False, indent=2)
        meta = {"dim": self.dim, "n_passages": len(self.passages)}
        with open(self.meta_path, "w", encoding="utf-8") as f:
            json.dump(meta, f, ensure_ascii=False, indent=2)
        logger.info("KB persisted: %s (passages=%d)", self.index_path, len(self.passages))

    def _load(self) -> None:
        if not self.index_path.exists() or not self.passages_path.exists():
            raise FileNotFoundError("Index or passages file missing")
        self.index = faiss.read_index(str(self.index_path))
        with open(self.passages_path, "r", encoding="utf-8") as f:
            self.passages = json.load(f)

    # -------------------------
    # Document ingestion
    # -------------------------
    def add_document(self, path: str) -> int:
        p = Path(path)
        if not p.exists():
            raise FileNotFoundError(f"Document not found: {path}")

        ext = p.suffix.lower()
        if ext == ".pdf":
            fulltext = _extract_text_from_pdf(p)
        elif ext in (".docx", ".doc"):
            fulltext = _extract_text_from_docx(p)
        elif ext == ".txt":
            fulltext = _read_txt(p)
        else:
            logger.warning("Skipping unsupported file type: %s", path)
            return 0

        if not fulltext.strip():
            logger.warning("No text extracted from %s — skipping", path)
            return 0

        chunks = _chunk_text(fulltext, chunk_size=self.chunk_size, overlap=self.chunk_overlap)
        embeddings = self.model.encode(chunks, convert_to_numpy=True, show_progress_bar=False).astype("float32")
        faiss.normalize_L2(embeddings)

        if self.index is None:
            self.index = faiss.IndexFlatIP(self.dim)
            logger.info("Created new FAISS index (dim=%d)", self.dim)

        self.index.add(embeddings)
        for c in chunks:
            self.passages.append({"text": c, "source": p.name})

        self._save()
        logger.info("Added %d chunks from %s", len(chunks), path)
        return len(chunks)

    def build_knowledge_base(self, sources_dir: Optional[str] = None) -> bool:
        candidates = []
        if sources_dir:
            candidates.append(Path(sources_dir))
        else:
            for p in ("Data Sources.pdf", "data", "resources"):
                candidates.append(Path(p))

        added_total = 0
        for c in candidates:
            if c.is_file():
                added_total += self.add_document(str(c))
            elif c.is_dir():
                for ext in ("*.pdf", "*.docx", "*.txt"):
                    for f in c.glob(ext):
                        added_total += self.add_document(str(f))

        return added_total > 0

    # -------------------------
    # Searching
    # -------------------------
    def search_knowledge(self, query: str, n_results: int = 3) -> Dict[str, Any]:
        if self.index is None or not self.passages:
            return {"query": query, "documents": []}

        q_vec = self.model.encode([query], convert_to_numpy=True).astype("float32")
        faiss.normalize_L2(q_vec)
        top_k = min(n_results, len(self.passages))
        distances, indices = self.index.search(q_vec, top_k)

        results = []
        for score, idx in zip(distances[0], indices[0]):
            if 0 <= idx < len(self.passages):
                results.append({
                    "text": self.passages[idx]["text"],
                    "source": self.passages[idx]["source"],
                    "score": float(score)
                })
        return {"query": query, "documents": results}

    def get_status(self) -> Dict[str, Any]:
        return {
            "status": "complete" if self.index and self.passages else "empty",
            "total_chunks": len(self.passages),
            "db_path": str(self.db_dir)
        }

    def clear_index(self) -> None:
        for path in (self.index_path, self.passages_path, self.meta_path):
            if path.exists():
                path.unlink()
        self.index = None
        self.passages = []
        logger.info("Cleared knowledge base.")


if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO)
    kb = ADGMKnowledgeBase()
    ok = kb.build_knowledge_base()
    print("KB build ok:", ok)
    print("Status:", kb.get_status())
    res = kb.search_knowledge("ADGM jurisdiction requirements", n_results=5)
    print(json.dumps(res, indent=2))
