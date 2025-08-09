from pathlib import Path
from typing import List, Dict, Any
import logging
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_COLOR_INDEX

logger = logging.getLogger(__name__)
logger.setLevel(logging.INFO)

class DocumentEditor:
    """
    Annotates DOCX files with highlights and a detailed appendix.
    The appendix now includes specific legal citations retrieved via RAG.
    """

    def __init__(self, output_dir: str = "annotated_docs"):
        self.out_dir = Path(output_dir)
        self.out_dir.mkdir(parents=True, exist_ok=True)

    def _highlight_run(self, run, severity: str):
        """Applies a highlight color based on issue severity."""
        try:
            color_map = {
                "CRITICAL": WD_COLOR_INDEX.RED,
                "HIGH": WD_COLOR_INDEX.ORANGE,
                "MEDIUM": WD_COLOR_INDEX.YELLOW,
                "LOW": WD_COLOR_INDEX.BRIGHT_GREEN,
            }
            highlight_color = color_map.get(severity.upper(), WD_COLOR_INDEX.AUTO)
            if highlight_color != WD_COLOR_INDEX.AUTO:
                run.font.highlight_color = highlight_color
                run.font.bold = True
        except Exception:
            # Fallback for environments where highlighting might fail
            run.font.bold = True
            run.font.underline = True

    def annotate_document(self, input_path: str, issues: List[Dict[str, Any]]) -> str:
        """
        Generates a reviewed .docx file with inline highlights and a formatted appendix
        containing suggestions and RAG-powered legal citations.
        """
        try:
            p = Path(input_path)
            doc = Document(str(p))

            # Normalize issues to ensure all keys are present
            norm_issues = []
            for i, issue_data in enumerate(issues):
                norm_issues.append({
                    "id": f"I{i+1}",
                    "issue": issue_data.get("issue") or issue_data.get("description", "No description provided."),
                    "suggestion": issue_data.get("suggestion", "No suggestion available."),
                    "severity": (issue_data.get("severity") or "MEDIUM").upper(),
                    "adgm_reference": issue_data.get("adgm_reference", ""),
                })

            # --- Inline Highlighting Logic ---
            # This part remains the same, finding and highlighting text.
            for para in doc.paragraphs:
                for issue in norm_issues:
                    # Use a simple heuristic to find the relevant phrase to highlight
                    phrase_to_find = issue["issue"].split(":")[-1].strip().lower()
                    if phrase_to_find and len(phrase_to_find) > 4 and phrase_to_find in para.text.lower():
                        for run in para.runs:
                            if phrase_to_find in run.text.lower():
                                self._highlight_run(run, issue["severity"])
                                run.add_text(f" [{issue['id']}]")

            # --- Appendix Generation (Updated for better formatting) ---
            doc.add_page_break()
            doc.add_heading("Review Comments (Automated Analysis)", level=1)
            doc.add_paragraph(
                "This document was reviewed by the ADGM Corporate Agent. Inline highlights correspond to the issues detailed below.",
                style='Intense Quote'
            )

            for issue in norm_issues:
                # Add the main issue description with its ID and severity
                p_entry = doc.add_paragraph(style="List Number")
                p_entry.add_run(f"[{issue['id']}] Issue: ").bold = True
                p_entry.add_run(f"{issue['issue']}")
                p_entry.add_run(f" (Severity: {issue['severity']})").italic = True

                # Add the suggestion, indented for clarity
                if issue["suggestion"]:
                    sugg_para = doc.add_paragraph()
                    sugg_para.paragraph_format.left_indent = Pt(36)
                    sugg_para.add_run("Suggestion: ").bold = True
                    sugg_para.add_run(issue["suggestion"])

                # **CRITICAL CHANGE**: Add the detailed ADGM reference from RAG, indented
                if issue["adgm_reference"]:
                    ref_para = doc.add_paragraph()
                    ref_para.paragraph_format.left_indent = Pt(36)
                    ref_para.add_run("ADGM Reference: ").bold = True
                    # The reference fetched by the validator is now added here
                    ref_para.add_run(f"{issue['adgm_reference']}").italic = True

            # Save the final document
            out_name = f"{p.stem}_reviewed.docx"
            out_path = self.out_dir / out_name
            doc.save(str(out_path))
            logger.info("Annotated document written: %s", out_path)
            return str(out_path)

        except Exception as e:
            logger.exception("Failed to annotate document: %s", e)
            raise